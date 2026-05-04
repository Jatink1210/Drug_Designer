"""Decision Dossier Generator — Requirements 9.1-9.5.

Assembles SynthArena session data into structured documents with sections:
Executive Summary, Compound Comparison, Scoring Matrix, Debate Summary,
Recommendation, Provenance Appendix.

Exports to PDF (WeasyPrint/HTML fallback) and DOCX (python-docx).
"""

import io
import json
import uuid
import zipfile
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

import structlog

log = structlog.get_logger(__name__)


class DossierContent:
    """Structured dossier content with all required sections."""

    def __init__(
        self,
        executive_summary: str = "",
        compound_comparison: Optional[Dict[str, Any]] = None,
        scoring_matrix: Optional[List[Dict[str, Any]]] = None,
        debate_summary: Optional[Dict[str, Any]] = None,
        recommendation: str = "",
        provenance_appendix: Optional[Dict[str, Any]] = None,
    ):
        self.executive_summary = executive_summary
        self.compound_comparison = compound_comparison or {}
        self.scoring_matrix = scoring_matrix or []
        self.debate_summary = debate_summary or {}
        self.recommendation = recommendation
        self.provenance_appendix = provenance_appendix or {
            "sources_consulted": [],
            "total_sources": 0,
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "method_used": "rule_based",
        }

    def to_dict(self) -> Dict[str, Any]:
        return {
            "executive_summary": self.executive_summary,
            "compound_comparison": self.compound_comparison,
            "scoring_matrix": self.scoring_matrix,
            "debate_summary": self.debate_summary,
            "recommendation": self.recommendation,
            "provenance_appendix": self.provenance_appendix,
        }


class DossierGenerator:
    """Generates decision dossiers from SynthArena session data."""

    async def generate(self, session_data: Dict[str, Any]) -> DossierContent:
        """Assemble session data into structured dossier sections.

        Requirements 9.1, 9.2: All required sections + provenance appendix.
        """
        session_id = session_data.get("session_id", str(uuid.uuid4())[:8])
        compounds = session_data.get("compounds", [])
        scoring = session_data.get("scoring", {})
        debate = session_data.get("debate", {})
        ranked = scoring.get("ranked_scenarios", []) or session_data.get("ranked_scenarios", [])

        # Track all sources consulted
        sources_consulted: List[Dict[str, Any]] = []

        # Executive Summary
        num_compounds = len(compounds)
        winner = debate.get("consensus", {}).get("winner_compound_id", "N/A")
        confidence = debate.get("consensus", {}).get("confidence", 0)
        exec_summary = (
            f"This dossier evaluates {num_compounds} compound(s) in SynthArena session {session_id}. "
        )
        if winner != "N/A":
            exec_summary += (
                f"The multi-agent debate concluded with {winner} as the recommended compound "
                f"(confidence: {confidence:.0%}). "
            )
        if ranked:
            best = ranked[0] if ranked else {}
            exec_summary += (
                f"Top-ranked scenario: {best.get('title', 'N/A')} "
                f"(composite score: {best.get('composite_score', 0):.3f})."
            )
        sources_consulted.append({"source": "SynthArena", "type": "session_data", "retrieved_at": datetime.now(timezone.utc).isoformat()})

        # Compound Comparison
        comparison = {
            "total_compounds": num_compounds,
            "compounds": [],
        }
        for c in compounds:
            cid = c.get("compound_id") or c.get("id") or c.get("name", "unknown")
            comparison["compounds"].append({
                "compound_id": cid,
                "name": c.get("name", cid),
                "smiles": c.get("smiles", ""),
                "properties": c.get("properties", {}),
                "scores": c.get("scores", {}),
            })
            sources_consulted.append({"source": f"compound_{cid}", "type": "compound_data", "retrieved_at": datetime.now(timezone.utc).isoformat()})

        # Scoring Matrix
        scoring_matrix = []
        for scenario in ranked:
            scoring_matrix.append({
                "scenario_id": scenario.get("scenario_id", ""),
                "title": scenario.get("title", ""),
                "composite_score": scenario.get("composite_score", 0),
                "signal_scores": scenario.get("signal_scores", {}),
                "risk_factors": scenario.get("risk_factors", []),
            })

        # Debate Summary
        debate_summary = {
            "agents": debate.get("agents", []),
            "total_arguments": len(debate.get("debate_history", [])),
            "consensus": debate.get("consensus", {}),
            "method_used": debate.get("method_used", "rule_based"),
        }

        # Recommendation
        if winner != "N/A" and confidence > 0.5:
            recommendation = (
                f"Based on multi-agent evaluation, compound {winner} is recommended for advancement. "
                f"Confidence level: {confidence:.0%}. "
            )
            dissenting = debate.get("consensus", {}).get("dissenting_opinions", [])
            if dissenting:
                recommendation += f"Note: {len(dissenting)} dissenting opinion(s) recorded."
        else:
            recommendation = "No clear recommendation. Further evaluation recommended."

        # Provenance Appendix
        provenance = {
            "sources_consulted": sources_consulted,
            "total_sources": len(sources_consulted),
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "method_used": debate.get("method_used", "rule_based"),
            "session_id": session_id,
        }

        return DossierContent(
            executive_summary=exec_summary,
            compound_comparison=comparison,
            scoring_matrix=scoring_matrix,
            debate_summary=debate_summary,
            recommendation=recommendation,
            provenance_appendix=provenance,
        )

    async def export_pdf(self, dossier: DossierContent) -> bytes:
        """Export dossier as PDF via WeasyPrint, with HTML fallback.

        Requirement 9.3.
        """
        html = self._render_html(dossier)
        try:
            from weasyprint import HTML
            pdf_bytes = HTML(string=html).write_pdf()
            return pdf_bytes
        except ImportError:
            log.warning("WeasyPrint not available, returning HTML as fallback")
            return html.encode("utf-8")
        except Exception as e:
            log.error("PDF generation failed: %s", e)
            return html.encode("utf-8")

    async def export_docx(self, dossier: DossierContent) -> bytes:
        """Export dossier as DOCX via python-docx.

        Requirement 9.4.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt

            doc = Document()
            doc.add_heading("Decision Dossier", level=0)

            # Executive Summary
            doc.add_heading("1. Executive Summary", level=1)
            doc.add_paragraph(dossier.executive_summary)

            # Compound Comparison
            doc.add_heading("2. Compound Comparison", level=1)
            compounds = dossier.compound_comparison.get("compounds", [])
            if compounds:
                table = doc.add_table(rows=1, cols=4)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "Compound"
                hdr[1].text = "SMILES"
                hdr[2].text = "Properties"
                hdr[3].text = "Scores"
                for c in compounds:
                    row = table.add_row().cells
                    row[0].text = str(c.get("name", ""))
                    row[1].text = str(c.get("smiles", ""))[:50]
                    row[2].text = json.dumps(c.get("properties", {}), indent=0)[:100]
                    row[3].text = json.dumps(c.get("scores", {}), indent=0)[:100]

            # Scoring Matrix
            doc.add_heading("3. Scoring Matrix", level=1)
            for entry in dossier.scoring_matrix:
                doc.add_paragraph(
                    f"{entry.get('title', 'N/A')}: {entry.get('composite_score', 0):.3f}",
                    style="List Bullet",
                )

            # Debate Summary
            doc.add_heading("4. Debate Summary", level=1)
            consensus = dossier.debate_summary.get("consensus", {})
            doc.add_paragraph(f"Winner: {consensus.get('winner_compound_id', 'N/A')}")
            doc.add_paragraph(f"Rationale: {consensus.get('winner_rationale', 'N/A')}")
            doc.add_paragraph(f"Confidence: {consensus.get('confidence', 0):.0%}")

            # Recommendation
            doc.add_heading("5. Recommendation", level=1)
            doc.add_paragraph(dossier.recommendation)

            # Provenance Appendix
            doc.add_heading("6. Provenance Appendix", level=1)
            prov = dossier.provenance_appendix
            doc.add_paragraph(f"Total sources consulted: {prov.get('total_sources', 0)}")
            doc.add_paragraph(f"Generated at: {prov.get('generated_at', '')}")
            for src in prov.get("sources_consulted", []):
                doc.add_paragraph(f"- {src.get('source', 'unknown')} ({src.get('type', '')})", style="List Bullet")

            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()

        except ImportError:
            log.warning("python-docx not available, returning JSON fallback")
            return json.dumps(dossier.to_dict(), indent=2).encode("utf-8")
        except Exception as e:
            log.error("DOCX generation failed: %s", e)
            return json.dumps(dossier.to_dict(), indent=2).encode("utf-8")

    def _render_html(self, dossier: DossierContent) -> str:
        """Render dossier as HTML for PDF generation."""
        consensus = dossier.debate_summary.get("consensus", {})
        compounds_html = ""
        for c in dossier.compound_comparison.get("compounds", []):
            compounds_html += f"<tr><td>{c.get('name', '')}</td><td><code>{c.get('smiles', '')[:40]}</code></td></tr>"

        scoring_html = ""
        for entry in dossier.scoring_matrix:
            scoring_html += f"<tr><td>{entry.get('title', '')}</td><td>{entry.get('composite_score', 0):.3f}</td></tr>"

        prov = dossier.provenance_appendix
        sources_html = ""
        for src in prov.get("sources_consulted", []):
            sources_html += f"<li>{src.get('source', 'unknown')} ({src.get('type', '')})</li>"

        return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>Decision Dossier</title>
<style>
body {{ font-family: Arial, sans-serif; margin: 40px; color: #1f2937; }}
h1 {{ color: #4338ca; border-bottom: 2px solid #4338ca; padding-bottom: 8px; }}
h2 {{ color: #1e40af; margin-top: 24px; }}
table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
th, td {{ border: 1px solid #d1d5db; padding: 8px 12px; text-align: left; }}
th {{ background: #f3f4f6; font-weight: 600; }}
.confidence {{ color: #059669; font-weight: bold; }}
code {{ background: #f3f4f6; padding: 2px 6px; border-radius: 4px; font-size: 0.85em; }}
</style></head><body>
<h1>Decision Dossier</h1>
<h2>1. Executive Summary</h2>
<p>{dossier.executive_summary}</p>
<h2>2. Compound Comparison</h2>
<table><tr><th>Compound</th><th>SMILES</th></tr>{compounds_html}</table>
<h2>3. Scoring Matrix</h2>
<table><tr><th>Scenario</th><th>Score</th></tr>{scoring_html}</table>
<h2>4. Debate Summary</h2>
<p><strong>Winner:</strong> {consensus.get('winner_compound_id', 'N/A')}</p>
<p><strong>Rationale:</strong> {consensus.get('winner_rationale', 'N/A')}</p>
<p><strong>Confidence:</strong> <span class="confidence">{consensus.get('confidence', 0):.0%}</span></p>
<h2>5. Recommendation</h2>
<p>{dossier.recommendation}</p>
<h2>6. Provenance Appendix</h2>
<p>Total sources: {prov.get('total_sources', 0)} | Generated: {prov.get('generated_at', '')}</p>
<ul>{sources_html}</ul>
</body></html>"""


# Legacy compatibility
class DossierCompiler:
    """Legacy dossier compiler for backward compatibility."""

    @staticmethod
    def generate_dossier_zip(project_data: dict) -> bytes:
        target_id = project_data.get("target_id", "undefined_target")
        log.info("dossier_compilation_started", target=target_id)

        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            metadata = {
                "compiled_at": datetime.now(timezone.utc).isoformat(),
                "agent_version": "Workbench Core v3.0",
                "target_id": target_id,
            }
            zf.writestr("metadata.json", json.dumps(metadata, indent=4))

            md_content = f"# Master Decision Dossier: {target_id}\n\n"
            md_content += f"**Compiled:** {metadata['compiled_at']}\n\n"
            md_content += "## 1. Executive Summary\n"
            md_content += project_data.get("llm_consensus", "No consensus generated.") + "\n\n"
            md_content += "## 2. Evidence\n"
            md_content += f"- **Sources analyzed:** {len(project_data.get('evidence_array', []))}\n"
            zf.writestr("executive_summary.md", md_content)

            evidence = project_data.get("evidence_array", [])
            zf.writestr("evidence_graph.json", json.dumps(evidence, indent=4))

            graph_topology = project_data.get("graph_topology", {})
            zf.writestr("viking_topology_nodes.json", json.dumps(graph_topology, indent=4))

        log.info("dossier_compilation_completed", size_bytes=buffer.tell())
        return buffer.getvalue()
