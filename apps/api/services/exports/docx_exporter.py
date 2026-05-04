"""DOCX export service for editable reports."""

from __future__ import annotations

from datetime import datetime
from typing import Any, Dict, List, Optional

import structlog
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

log = structlog.get_logger()


class DOCXReportExporter:
    """
    DOCX exporter for editable scientific reports.
    
    Generates Microsoft Word-compatible documents that users can edit.
    Compatible with Microsoft Word 2016+.
    """
    
    def __init__(
        self,
        output_path: str,
        title: str,
        author: str = "Drug Designer System"
    ):
        """
        Initialize DOCX exporter.
        
        Args:
            output_path: Path to output DOCX file
            title: Document title
            author: Document author
        """
        self.output_path = output_path
        self.title = title
        self.author = author
        
        # Create document
        self.doc = Document()
        
        # Set document properties
        self.doc.core_properties.title = title
        self.doc.core_properties.author = author
        self.doc.core_properties.created = datetime.now()
        
        log.info("docx_exporter_initialized", output_path=output_path)
    
    def add_title_page(
        self,
        subtitle: Optional[str] = None,
        date: Optional[str] = None
    ) -> None:
        """Add title page to document."""
        # Title
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(self.title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(26, 26, 26)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        self.doc.add_paragraph()
        
        # Subtitle
        if subtitle:
            subtitle_para = self.doc.add_paragraph()
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.size = Pt(16)
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.doc.add_paragraph()
        
        # Author and date
        author_para = self.doc.add_paragraph()
        author_run = author_para.add_run(f"Author: {self.author}")
        author_run.font.size = Pt(12)
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_str = date or datetime.now().strftime("%B %d, %Y")
        date_para = self.doc.add_paragraph()
        date_run = date_para.add_run(f"Date: {date_str}")
        date_run.font.size = Pt(12)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Page break
        self.doc.add_page_break()
    
    def add_heading(
        self,
        text: str,
        level: int = 1
    ) -> None:
        """
        Add a heading to the document.
        
        Args:
            text: Heading text
            level: Heading level (1-3)
        """
        self.doc.add_heading(text, level=level)
    
    def add_paragraph(
        self,
        text: str,
        bold: bool = False,
        italic: bool = False
    ) -> None:
        """
        Add a paragraph to the document.
        
        Args:
            text: Paragraph text
            bold: Make text bold
            italic: Make text italic
        """
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(11)
        
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
    
    def add_table(
        self,
        data: List[List[str]],
        has_header: bool = True
    ) -> None:
        """
        Add a table to the document.
        
        Args:
            data: Table data (list of rows)
            has_header: Whether first row is header
        """
        if not data:
            return
        
        # Create table
        table = self.doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Light Grid Accent 1'
        
        # Populate table
        for i, row_data in enumerate(data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_data)
                
                # Style header row
                if i == 0 and has_header:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    cell._element.get_or_add_tcPr().append(
                        cell._element._new_tblPr()
                    )
        
        self.doc.add_paragraph()  # Add spacing after table
    
    def add_bullet_list(
        self,
        items: List[str]
    ) -> None:
        """
        Add a bullet list to the document.
        
        Args:
            items: List items
        """
        for item in items:
            self.doc.add_paragraph(item, style='List Bullet')
    
    def add_numbered_list(
        self,
        items: List[str]
    ) -> None:
        """
        Add a numbered list to the document.
        
        Args:
            items: List items
        """
        for item in items:
            self.doc.add_paragraph(item, style='List Number')
    
    def add_section(
        self,
        title: str,
        content: str,
        level: int = 1
    ) -> None:
        """
        Add a section to the document.
        
        Args:
            title: Section title
            content: Section content
            level: Heading level
        """
        self.add_heading(title, level=level)
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                self.add_paragraph(para.strip())
    
    def export_report(
        self,
        report_data: Dict[str, Any]
    ) -> str:
        """
        Export complete report to DOCX.
        
        Args:
            report_data: Report data including sections, tables, and lists
            
        Returns:
            Path to generated DOCX file
        """
        # Add title page
        self.add_title_page(
            subtitle=report_data.get('subtitle'),
            date=report_data.get('date')
        )
        
        # Add sections
        for section in report_data.get('sections', []):
            self.add_section(
                title=section.get('title', ''),
                content=section.get('content', ''),
                level=section.get('level', 1)
            )
            
            # Add tables if present
            for table in section.get('tables', []):
                self.add_table(
                    data=table.get('data', []),
                    has_header=table.get('has_header', True)
                )
            
            # Add lists if present
            for bullet_list in section.get('bullet_lists', []):
                self.add_bullet_list(bullet_list)
            
            for numbered_list in section.get('numbered_lists', []):
                self.add_numbered_list(numbered_list)
        
        # Save document
        try:
            self.doc.save(self.output_path)
            log.info("docx_export_complete", output_path=self.output_path)
            return self.output_path
        except Exception as e:
            log.error("docx_export_failed", error=str(e))
            raise
